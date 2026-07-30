#!/usr/bin/env python3
"""
doc_extract.py — конвертація документів на товар у ЧИСТИЙ ТЕКСТ (+ Markdown)
та детерміноване витягування структурованих ідентифікаторів, які дають
«буст точності» аналізу (CAS №, UN №, коди HS/УКТЗЕД, чистота, виробник,
країна походження, партія).

Навіщо: у SDS/CoA/інвойсах є ОДНОЗНАЧНІ ідентифікатори (CAS, UN, HS), надійніші
за розбір вільної назви з маніфесту. CAS → речовина → код; UN → клас ADR;
явний HS → авторитетний код. Скрипт готує ці сигнали + чистий текст для AI.

Формати: PDF (текст; скани — через OCR, якщо встановлено tesseract), DOCX,
XLSX/XLS, зображення (PNG/JPG/TIFF — OCR), TXT/CSV/MD.

Витягування: ДЕТЕРМІНОВАНЕ (regex + валідація: контрольна сума CAS, перевірка
HS-6 проти реальної номенклатури) + опційний AI-fallback (--ai) для полів,
яких не знайдено.

Використання:
  python scripts/doc_extract.py <файл|тека> [-o out_dir] [--ai] [--lang eng+chi_sim+ukr]

Вивід: для кожного файла out_dir/<name>.md (чистий текст) і out_dir/<name>.json
(поля). Зведення друкується у stdout (JSON).

Залежності (базові вже є): PyMuPDF (fitz), python-docx, openpyxl, Pillow.
OCR (опційно, для сканів): pip install pytesseract  +  системний tesseract
(https://github.com/UB-Mannheim/tesseract). Мовні пакети: eng, chi_sim, ukr.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

# UTF-8 у консоль (щоб кирилиця в stderr не билася на Windows-кодуванні).
for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8")
    except Exception:
        pass

# ── Опційні залежності (грейсфул-деградація) ─────────────────────────────
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    import docx  # python-docx
except Exception:
    docx = None
try:
    import openpyxl
except Exception:
    openpyxl = None
try:
    from PIL import Image
except Exception:
    Image = None
try:
    import pytesseract
except Exception:
    pytesseract = None

REPO = Path(__file__).resolve().parent.parent

# Реальні 6-значні підпозиції HS — для валідації кандидатів кодів (висока точність).
try:
    HS_VALID6 = set(json.loads((REPO / "lib" / "data" / "hs_valid6.json").read_text(encoding="utf-8")))
except Exception:
    HS_VALID6 = set()


# ── Очищення тексту ──────────────────────────────────────────────────────
def clean_text(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"-\n(?=\w)", "", s)          # склеювання переносів по дефісу
    s = re.sub(r"\n{3,}", "\n\n", s)          # не більше одного порожнього рядка
    return s.strip()


# ── Конвертація за типом файла ───────────────────────────────────────────
def ocr_available() -> bool:
    if pytesseract is None:
        return False
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def ocr_image(img, lang: str) -> str:
    if not ocr_available():
        return ""
    try:
        return pytesseract.image_to_string(img, lang=lang)
    except Exception as e:
        print(f"  [OCR] помилка: {e}", file=sys.stderr)
        return ""


def pdf_to_text(path: Path, lang: str, meta: dict) -> str:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) не встановлено — потрібен для PDF.")
    out = []
    doc = fitz.open(path)
    for page in doc:
        txt = page.get_text("text") or ""
        # Сторінка-скан (мало/нема тексту) → пробуємо OCR по рендеру.
        if len(txt.strip()) < 20:
            if ocr_available() and Image is not None:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                ocr = ocr_image(img, lang)
                if ocr.strip():
                    meta["ocr_used"] = True
                    txt = ocr
                else:
                    meta["needs_ocr"] = True
            else:
                meta["needs_ocr"] = True
        out.append(txt)
    doc.close()
    return "\n\n".join(out)


def docx_to_text(path: Path) -> str:
    if docx is None:
        raise RuntimeError("python-docx не встановлено.")
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def xlsx_to_text(path: Path) -> str:
    if openpyxl is None:
        raise RuntimeError("openpyxl не встановлено.")
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def image_to_text(path: Path, lang: str, meta: dict) -> str:
    if Image is None:
        raise RuntimeError("Pillow не встановлено.")
    img = Image.open(path)
    txt = ocr_image(img, lang)
    if txt.strip():
        meta["ocr_used"] = True
    else:
        meta["needs_ocr"] = True
    return txt


def convert(path: Path, lang: str) -> tuple[str, dict]:
    meta = {"ocr_used": False, "needs_ocr": False}
    ext = path.suffix.lower()
    if ext == ".pdf":
        raw = pdf_to_text(path, lang, meta)
    elif ext in (".docx",):
        raw = docx_to_text(path)
    elif ext in (".xlsx", ".xlsm"):
        raw = xlsx_to_text(path)
    elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"):
        raw = image_to_text(path, lang, meta)
    elif ext in (".txt", ".md", ".csv"):
        raw = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise RuntimeError(f"Непідтримуваний формат: {ext}")
    return clean_text(raw), meta


# ── Детерміновані екстрактори ────────────────────────────────────────────
def validate_cas(cas: str) -> bool:
    """CAS = N..N-NN-C; C — контрольна цифра (Σ(цифра*позиція справа) mod 10)."""
    digits = cas.replace("-", "")
    if len(digits) < 5:
        return False
    body, check = digits[:-1], int(digits[-1])
    s = sum(int(d) * (i + 1) for i, d in enumerate(reversed(body)))
    return s % 10 == check


def extract_cas(text: str) -> list[str]:
    found = []
    for m in re.finditer(r"\b(\d{2,7}-\d{2}-\d)\b", text):
        cas = m.group(1)
        if validate_cas(cas) and cas not in found:
            found.append(cas)
    return found


def extract_un(text: str) -> list[str]:
    found = []
    for m in re.finditer(r"\bUN[\s:./-]?(\d{4})\b", text, re.IGNORECASE):
        un = m.group(1)
        if 1 <= int(un) <= 3600 and un not in found:
            found.append(un)
    return found


def extract_hs(text: str) -> list[str]:
    """10-значні коди + 6/8-значні біля ключових слів; валідація HS-6."""
    cands: list[str] = []
    # 10 знаків (можливо з пробілами/крапками)
    for m in re.finditer(r"(?<!\d)(\d{4}[\s.]?\d{2}[\s.]?\d{2}[\s.]?\d{2})(?!\d)", text):
        d = re.sub(r"\D", "", m.group(1))
        if len(d) == 10:
            cands.append(d)
    # 6/8 знаків поряд із HS/УКТЗЕД/код/tariff
    for m in re.finditer(
        r"(?:hs|hs\s*code|уктзед|тнзед|код|tariff|н\.?в\.?е\.?д)[^\d]{0,12}(\d{4}[\s.]?\d{2}(?:[\s.]?\d{2})?)",
        text,
        re.IGNORECASE,
    ):
        d = re.sub(r"\D", "", m.group(1))
        if len(d) in (6, 8):
            cands.append(d)
    # Валідація HS-6 (якщо словник наявний). Пусті знання → не фільтруємо.
    valid = []
    for c in cands:
        if HS_VALID6 and c[:6] not in HS_VALID6:
            continue
        if c not in valid:
            valid.append(c)
    # Дедуп: прибираємо коротший код, який є префіксом довшого (10 знак > 8 > 6).
    valid.sort(key=len, reverse=True)
    out: list[str] = []
    for c in valid:
        if any(longer.startswith(c) for longer in out):
            continue
        out.append(c)
    return out


def _find_after(text: str, keys: str) -> str | None:
    m = re.search(rf"(?:{keys})\s*[:\-–]\s*([^\n]{{2,80}})", text, re.IGNORECASE)
    return m.group(1).strip(" .;,") if m else None


def extract_purity(text: str) -> list[str]:
    out = []
    for m in re.finditer(
        r"(?:purity|assay|content|чистота|вміст|вмiст|содержание|min\.?)\D{0,12}(\d{2,3}(?:[.,]\d+)?)\s*%",
        text,
        re.IGNORECASE,
    ):
        v = m.group(1).replace(",", ".")
        if v not in out:
            out.append(v + "%")
    return out


def extract_fields(text: str) -> dict:
    return {
        "cas": extract_cas(text),
        "un": extract_un(text),
        "hs": extract_hs(text),
        "purity": extract_purity(text),
        "manufacturer": _find_after(text, r"manufacturer|виробник|производитель|mfr|made by|producer"),
        "country": _find_after(text, r"country of origin|країна походження|страна происхождения|made in|origin"),
        "batch": _find_after(text, r"batch(?:\s*(?:no|№|number))?|lot(?:\s*no)?|партія|серія|серия"),
        "product": _find_after(text, r"product name|назва товару|наименование|substance|product"),
    }


def missing_fields(fields: dict) -> list[str]:
    empty = []
    for k, v in fields.items():
        if v is None or (isinstance(v, list) and not v):
            empty.append(k)
    return empty


# ── Опційний AI-fallback (OpenAI-сумісний endpoint) ──────────────────────
def ai_fill(text: str, want: list[str], max_chars: int) -> dict:
    """Добирає відсутні поля через AI. Потрібні env: DOC_AI_BASE_URL, DOC_AI_KEY, DOC_AI_MODEL."""
    import urllib.request

    base = os.environ.get("DOC_AI_BASE_URL", "https://api.openai.com/v1")
    key = os.environ.get("DOC_AI_KEY", "")
    model = os.environ.get("DOC_AI_MODEL", "gpt-4o-mini")
    if not key:
        print("  [AI] пропуск: не задано DOC_AI_KEY.", file=sys.stderr)
        return {}
    prompt = (
        "Витягни з тексту документа на товар лише ці поля JSON: "
        + ", ".join(want)
        + ". cas/un/hs/purity — масиви рядків; manufacturer/country/batch/product — рядок або null. "
        "Не вигадуй значень, яких немає в тексті. Поверни СТРОГИЙ JSON без markdown.\n\nТЕКСТ:\n"
        + text[:max_chars]
    )
    body = json.dumps(
        {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0,
            "response_format": {"type": "json_object"},
        }
    ).encode("utf-8")
    req = urllib.request.Request(
        f"{base}/chat/completions",
        data=body,
        headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            data = json.loads(r.read())
        content = data["choices"][0]["message"]["content"]
        return json.loads(content)
    except Exception as e:
        print(f"  [AI] помилка: {e}", file=sys.stderr)
        return {}


# ── Оркестрація ──────────────────────────────────────────────────────────
SUPPORTED = {".pdf", ".docx", ".xlsx", ".xlsm", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".txt", ".md", ".csv"}


def process_file(path: Path, out_dir: Path, lang: str, use_ai: bool, max_chars: int) -> dict:
    text, meta = convert(path, lang)
    fields = extract_fields(text)

    if use_ai:
        want = missing_fields(fields)
        if want and text.strip():
            filled = ai_fill(text, want, max_chars)
            for k in want:
                if k in filled and filled[k]:
                    fields[k] = filled[k]
                    meta.setdefault("ai_filled", []).append(k)

    stem = path.stem
    (out_dir / f"{stem}.md").write_text(text, encoding="utf-8")
    result = {"file": path.name, "chars": len(text), **meta, "fields": fields}
    (out_dir / f"{stem}.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    flag = " [OCR]" if meta.get("ocr_used") else (" [ПОТРІБЕН OCR]" if meta.get("needs_ocr") else "")
    hs = ",".join(fields["hs"]) or "—"
    cas = ",".join(fields["cas"]) or "—"
    print(f"✓ {path.name}{flag}  chars={len(text)}  CAS={cas}  HS={hs}", file=sys.stderr)
    return result


def main() -> int:
    ap = argparse.ArgumentParser(description="Конвертація документів → чистий текст + витяг ідентифікаторів.")
    ap.add_argument("path", help="Файл або тека з документами.")
    ap.add_argument("-o", "--out", default="doc_out", help="Тека виводу (default: doc_out).")
    ap.add_argument("--ai", action="store_true", help="Добирати відсутні поля через AI-fallback.")
    ap.add_argument("--lang", default="eng", help="Мови OCR (напр. eng+chi_sim+ukr).")
    ap.add_argument("--max-chars", type=int, default=12000, help="Ліміт тексту для AI.")
    args = ap.parse_args()

    src = Path(args.path)
    if not src.exists():
        print(f"Немає шляху: {src}", file=sys.stderr)
        return 1
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    files = (
        [p for p in sorted(src.rglob("*")) if p.suffix.lower() in SUPPORTED]
        if src.is_dir()
        else [src]
    )
    if not files:
        print("Не знайдено підтримуваних файлів.", file=sys.stderr)
        return 1

    if not ocr_available():
        print("  [i] OCR недоступний (нема pytesseract/tesseract) — скани не читатимуться. "
              "Встанови: pip install pytesseract + системний tesseract.", file=sys.stderr)

    results = []
    for f in files:
        try:
            results.append(process_file(f, out_dir, args.lang, args.ai, args.max_chars))
        except Exception as e:
            print(f"✗ {f.name}: {e}", file=sys.stderr)
            results.append({"file": f.name, "error": str(e)})

    print(json.dumps({"out_dir": str(out_dir), "count": len(results), "results": results}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
